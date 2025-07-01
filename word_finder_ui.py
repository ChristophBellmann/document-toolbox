import os
import re
import requests
from collections import Counter
from tkinter import Frame, Button, Label, Text, Entry, Scrollbar, END, messagebox

from docx import Document as DocxDocument

word_folder = os.path.join(os.getcwd(), "Word")
fuellwort_datei = os.path.join(word_folder, "Fuellwoerter.txt")

def load_fuellwoerter():
    if os.path.exists(fuellwort_datei):
        with open(fuellwort_datei, "r") as f:
            return set(line.strip().lower() for line in f)
    return set()

def add_fuellwort(word):
    with open(fuellwort_datei, "a") as f:
        f.write(word + "\n")
    fuellwoerter.add(word)

def replace_word_in_docx(file_path, old_word, new_word):
    doc = DocxDocument(file_path)
    pattern = re.compile(r'\b' + re.escape(old_word) + r'\b', flags=re.IGNORECASE)

    matches = []
    for i, para in enumerate(doc.paragraphs):
        for m in pattern.finditer(para.text):
            matches.append((i, m.start(), m.end()))

    if not matches:
        messagebox.showinfo("Info", "Wort nicht gefunden.")
        return

    last_para_idx, start, end = matches[-1]
    para = doc.paragraphs[last_para_idx]

    text = para.text
    new_text = text[:start] + new_word + text[end:]
    para.text = new_text

    doc.save(file_path)
    messagebox.showinfo("Erfolg", f"Ersetzt im Absatz {last_para_idx}: '{old_word}' durch '{new_word}'.")

def show_sentence_with_word(file_path, word, text_widget):
    doc = DocxDocument(file_path)
    for para in doc.paragraphs:
        if word in para.text.lower():
            text_widget.delete(1.0, END)
            text_widget.insert(END, para.text.strip())
            break

def show_synonyms(word, text_widget):
    text_widget.delete(1.0, END)
    text_widget.insert(END, f"Synonym-Vorschläge für '{word}':\n")
    try:
        response = requests.get(f"https://www.openthesaurus.de/synonyme/search?q={word}&format=application/json")
        data = response.json()
        synsets = data.get("synsets", [])
        synonyms = []
        for synset in synsets:
            synonyms.extend(term["term"] for term in synset.get("terms", []))
        if synonyms:
            filename = doc_files[current_index]
            file_path = os.path.join(word_folder, filename)
            doc_text = " ".join([para.text.lower() for para in DocxDocument(file_path).paragraphs])
            count = 0
            for s in synonyms:
                pattern = r'\b' + re.escape(s.lower()) + r'\b'
                if not re.search(pattern, doc_text):
                    text_widget.insert(END, "- " + s + "\n")
                    count += 1
                if count == 3:
                    break
            if count == 0:
                text_widget.insert(END, "Keine neuen Synonyme gefunden.\n")
        else:
            text_widget.insert(END, "Keine Synonyme gefunden.\n")
    except Exception as e:
        text_widget.insert(END, f"Fehler beim Abrufen von Synonymen: {e}\n")

def load_doc_words():
    global word_buttons
    word_buttons.clear()
    for widget in word_frame.winfo_children():
        widget.destroy()

    if not doc_files:
        output_text.delete(1.0, END)
        output_text.insert(END, "Keine Dokumente gefunden.\n")
        return

    filename = doc_files[current_index]
    file_path = os.path.join(word_folder, filename)
    output_text.delete(1.0, END)
    output_text.insert(END, f"Dokument: {filename}\n")

    doc_text = " ".join([para.text for para in DocxDocument(file_path).paragraphs])
    words = [word.lower() for word in doc_text.split() if word.isalpha()]
    word_counts = Counter(words)

    for word, count in word_counts.items():
        if count > 1 and word not in fuellwoerter:
            frame = Frame(word_frame)
            frame.pack(fill="x", pady=2)

            btn_add = Button(frame, text=f"{word}: {count}", width=15, command=lambda w=word, f=frame: mark_fuellwort(w, f))
            btn_add.pack(side="left")

            btn_replace = Button(frame, text="✎", fg="red", command=lambda w=word, f=frame: replace_word_ui(w, f))
            btn_replace.pack(side="left")

            word_buttons.append(frame)

def mark_fuellwort(word, frame):
    add_fuellwort(word)
    frame.destroy()

def replace_word_ui(word, frame):
    for widget in frame.winfo_children():
        widget.destroy()

    Label(frame, text=f"Wort: {word}").pack(side="left")

    entry = Entry(frame)
    entry.pack(side="left")

    def on_confirm():
        new_word = entry.get().strip()
        if new_word:
            filename = doc_files[current_index]
            file_path = os.path.join(word_folder, filename)
            replace_word_in_docx(file_path, word, new_word)
            frame.destroy()
            load_doc_words()

    Button(frame, text="Bestätigen", command=on_confirm).pack(side="left")

    def on_show_sentence():
        filename = doc_files[current_index]
        file_path = os.path.join(word_folder, filename)
        show_sentence_with_word(file_path, word, output_text)

    def on_show_synonyms():
        show_synonyms(word, output_text)

    Button(frame, text="Satz anzeigen", command=on_show_sentence).pack(side="left")
    Button(frame, text="Synonyme", command=on_show_synonyms).pack(side="left")

def prev_doc():
    global current_index
    current_index = (current_index - 1) % len(doc_files)
    load_doc_words()

def next_doc():
    global current_index
    current_index = (current_index + 1) % len(doc_files)
    load_doc_words()

# Hauptfunktion, die vom Hauptfenster aufgerufen wird
def create_word_finder_ui(parent):
    global word_frame, output_text, fuellwoerter, doc_files, current_index, word_buttons

    fuellwoerter = load_fuellwoerter()
    doc_files = [f for f in os.listdir(word_folder) if f.endswith(".docx")]
    current_index = 0
    word_buttons = []

    nav_frame = Frame(parent)
    nav_frame.pack(pady=5)

    Button(nav_frame, text="<< Vorheriges Dokument", command=prev_doc).pack(side="left", padx=5)
    Button(nav_frame, text="Nächstes Dokument >>", command=next_doc).pack(side="left", padx=5)

    word_frame = Frame(parent)
    word_frame.pack(fill="both", expand=True, pady=10)

    output_text = Text(parent, height=10, wrap="word")
    output_text.pack(fill="both", expand=True, padx=10, pady=5)

    scrollbar = Scrollbar(output_text)
    scrollbar.pack(side="right", fill="y")

    load_doc_words()
